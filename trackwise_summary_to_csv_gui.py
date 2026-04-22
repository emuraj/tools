#!/usr/bin/env python3
"""
trackwise_summary_to_csv_gui_realtime.py

REAL-TIME writer version:
- Creates output files immediately (CSV + event log)
- Writes one row per record as it completes
- Flushes periodically
- Still prints RUN SUMMARY + full ERROR LIST at end
"""

import os
import re
import sys
import csv
import time
import statistics
import threading
import queue
from dataclasses import dataclass, asdict
from datetime import datetime
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

from docx import Document

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

try:
    import win32com.client  # type: ignore
    HAS_WIN32 = True
except Exception:
    HAS_WIN32 = False


# -----------------------------
# Config
# -----------------------------

TITLE_CHAR_LIMIT = 95
DEFAULT_MODEL = "gpt-4.1"

FLUSH_EVERY_N_RECORDS = 10  # flush CSVs + event log every N written records

DISCLAIMER_TEXT = (
    "Legacy Migrated Record Disclaimer: This file represents a reconstructed digital record derived from a "
    "TrackWise® Quality Management System export. TrackWise stores GMP records across many inter-related .csv "
    "files and associated binary attachments, with information distributed across tables using cryptic system-"
    "generated identifiers. The contents of this reconstructed record were generated through a deterministic "
    "and fully auditable algorithmic process that parsed relevant TrackWise source tables, resolved foreign-key "
    "relationships, reconstructed the hierarchical record structure, retrieved associated original file "
    "attachments (“blobs”), produced a structured JSON representation, and generated a DOCX summary and compiled "
    "the reconstructed summary plus associated files into a single, paginated final PDF suitable for inspection "
    "and long-term archival. The “Additional Files” provided alongside this record are the standalone source "
    "components used in compilation, including the reconstructed JSON, the DOCX summary, and the original "
    "attachments; these items are provided individually for traceability and—after conversion where applicable—"
    "also comprise the merged final PDF. No information was inferred, altered, or omitted. CR-00147 is the "
    "governing Change Control that authorizes and defines the migration/reconstruction of legacy TrackWise "
    "records into the new QMS record format and archival package."
)


# -----------------------------
# Prompt engineering (base + addenda + full examples)
# -----------------------------

BASE_PROMPT = f"""
PURPOSE:
Generate a professional GMP “prep brief” for a legacy TrackWise record. It must be auditor-friendly and teammate-prep-friendly.

STRICT OUTPUT (ONLY, in this order):
1) Title (one line). Must be <= {TITLE_CHAR_LIMIT} characters. Format: RECORD-ID – concise keyworded topic.
2) Prep Brief Synopsis (1–2 short paragraphs). Closure-focused and human/professional (not database-field narration).
3) {DISCLAIMER_TEXT}

TITLE RULES:
- Hard cap <= {TITLE_CHAR_LIMIT} chars. If too long, rewrite shorter (do not truncate mid-word).
- High-signal keywords; avoid long lists.

DATE RULES (GMP FORMAT):
- All dates must be DDMMMYY.
- Closure-first rule:
  - If the record is closed, state “closed DDMMMYY” (do NOT say “last recorded activity”).
  - If not closed, state “last recorded activity DDMMMYY”.
- Include approval/effective dates only if they are central evidence anchors.

SYNOPSIS MUST COVER:
- Why the record exists (trigger/context; include linked IDs only if explicitly present)
- What was done / what changed (material actions only)
- Closure basis (effectiveness required vs not required, with recorded rationale)
- Primary evidence anchors (objective artifacts proving closure)

CRITICAL AVOID:
- No citations, no URLs, no reference markers, no footnotes.
- No TrackWise workflow trivia (“final workflow action…”, “status field…”, etc.)
- No speculation; omit anything not explicit.
- Do NOT mention CR-00147 anywhere except within the disclaimer text in section 3.
- Do not output anything beyond the three required sections.

LENGTH TARGET:
- Synopsis should be ~120–220 words unless the record requires slightly more for clarity.

STYLE EXEMPLARS:
- The following exemplar is provided ONLY to demonstrate tone/structure.
- Do NOT reuse exemplar facts. Use ONLY the provided SOURCE_TEXT.
""".strip()

CAPA_ADDENDUM = """
CAPA-SPECIFIC:
Write as a CAPA prep brief. State the initiating driver (DEV/audit/observation) only if explicitly present.
Summarize the corrective/preventive actions implemented and closure basis. If an EC is referenced, include the EC ID.
Emphasize objective evidence anchors: approved SOP/form revisions, CMMS schedule changes, training completion, retest results, executed agreements, etc.
Avoid procedural history and approvals mechanics.
""".strip()

CR_ADDENDUM = """
CR-SPECIFIC:
Write as a change control prep brief. Summarize purpose/scope of the change, key impact controls (validation, training, doc updates),
implementation verification, and closure evidence. Emphasize evidence anchors rather than workflow routing.
""".strip()

DEV_ADDENDUM = """
DEV-SPECIFIC:
Write as a deviation prep brief. Include:
(1) what occurred and where (room/system/operation),
(2) occurrence date if present and record status using closure-first rule,
(3) deviation classification (Minor/Major/Critical) if present, and recorded impact statement (explicitly whether SISPQ/patient/process/product impacted),
(4) immediate actions/containment,
(5) investigation conclusion (root cause category + succinct rationale as recorded),
(6) CAPA determination (CAPA required yes/no with recorded justification).

TITLE REQUIREMENT FOR DEV:
If Minor/Major/Critical is explicitly stated in the record, include it in the title (e.g., “(Major)”).

Primary evidence anchors: excursion packet/results, cleaning actions, batch review, retest/resample, and recurrence/trending review as applicable.
Do not include long 6M narratives, audit trails, or distribution lists.
""".strip()

EC_ADDENDUM = """
EC-SPECIFIC:
Write as an effectiveness-check prep brief. First determine whether a recorded conclusion/result exists.
- If no conclusion is present: treat as an approved plan (not executed) and do NOT claim effectiveness.
- If concluded: summarize the conclusion and the evidence basis briefly.

Include: objective, related records explicitly listed, timing logic (e.g., months post-closure) and due date if present,
evaluation criteria/success checks as recorded, and evidence that will demonstrate completion (training, logs/work orders, config evidence, lookback/trending).
Apply closure-first rule for closed ECs; otherwise state last recorded activity.
""".strip()

CAPA_EXAMPLE = """
STYLE EXEMPLAR (CAPA) — DO NOT REUSE FACTS:

CAPA-00XXX – Data Integrity Controls: Unique Logins, Admin Segregation, and Role Review

CAPA-00XXX is a legacy TrackWise CAPA opened on 11DEC23 and closed 02JUL24 to remediate identified Data Integrity vulnerabilities in GxP computerized systems related to user access control and administrative privilege use. It answers the audit question of whether system access is uniquely assigned, least-privilege aligned, and supported by objective review evidence that prevents shared credentials and improper use of administrative accounts.

The record documents a system-by-system access assessment across the defined GxP system inventory, with remediation actions to disable or control generic/shared accounts, confirm unique named-user access, and restrict administrative privileges to appropriate personnel with segregation from routine operator activities. Where system constraints require dual-role capability, the record documents the justification and the compensating control expectations. Closure is supported by completed access review artifacts for each system in scope (as-found vs. as-left), updated role/permission definitions where needed, and documented evidence that corrective actions were implemented in the source systems. Effectiveness verification is recorded as Not Required, consistent with a governance correction supported by objective access review evidence and sustained administrative control expectations.
""".strip()

CR_EXAMPLE = """
STYLE EXEMPLAR (CR) — DO NOT REUSE FACTS:

CR-00XXX – Equipment/Facility Change Implementation: Validation Evidence, SOP Updates, and Training

CR-00XXX is a legacy TrackWise change control opened on 03JAN24 and closed 15APR24 to implement a defined equipment/facility/process change under controlled risk assessment and verification expectations. It answers the audit question of whether the change was evaluated for GMP impact, implemented as approved, and verified through objective evidence (validation/testing, documentation updates, and training) sufficient to support continued controlled operations.

The record documents the approved change scope, required impact controls, and implementation deliverables, including updates to affected controlled documents (SOPs/forms), execution of required validation or qualification activities, and completion of training for impacted roles prior to use. Where supplier or material/service dependencies exist, the record documents how readiness was established (e.g., supplier qualification evidence, executed agreements, or receipt/acceptance documentation). Closure is supported by the completed implementation package showing that verification activities met acceptance criteria, required documents were updated and effective, and training was assigned/completed in alignment with the change.
""".strip()

DEV_EXAMPLE = """
STYLE EXEMPLAR (DEV) — DO NOT REUSE FACTS:

DEV-00XXX (Major) – ISO 7 EM Action Limit Excursion: Mold Recovery, Containment, No Product Impact

DEV-00XXX is a legacy TrackWise deviation opened on 12MAR25 and closed 20MAR25 to document an ISO 7 environmental monitoring action limit excursion driven by a mold recovery at a defined location during routine viable EM. It answers the audit question of whether the excursion was promptly contained, assessed for impact to product/patient and process suitability (SISPQ), investigated to an appropriate depth, and closed with a defensible CAPA determination.

The record documents immediate containment actions including targeted area cleaning with sporicidal disinfectant, execution of follow-up monitoring at the affected location(s), and confirmation that operations were controlled while the assessment was completed. The impact assessment is recorded as no impact to product/patient and no confirmed impact to process suitability, with disposition based on the documented review of operations and the follow-up EM results. The investigation conclusion is documented as an environmental/housekeeping control variability cause category (as recorded), supported by site conditions and the absence of additional adverse signals during follow-up. CAPA determination is recorded as No CAPA required, with justification that the event was isolated, containment was effective, and there was no recurrence signal within the defined follow-up window. Primary evidence includes the EM excursion packet, documented cleaning/sporicidal execution, and follow-up EM demonstrating return to baseline.
""".strip()

EC_EXAMPLE = """
STYLE EXEMPLAR (EC) — DO NOT REUSE FACTS:

EC-00XXX – 12-Month Effectiveness Review Plan: Sustained Implementation and No Recurrence

EC-00XXX is a legacy TrackWise effectiveness check opened on 05APR25 to verify sustained effectiveness of the corrective actions implemented under the associated CAPA(s) and/or deviation(s) referenced in the record. This record is documented as an approved plan (no recorded conclusion present). It answers the audit question of whether the implemented controls remain in place over time, are used as intended, and prevent recurrence of the original failure mode after the initial corrective actions were completed.

The plan specifies the review timing relative to the associated record closure and includes a due date of 05APR26. Effectiveness criteria are defined as objective checks, including confirmation that revised procedures/system controls remain current and in use, required training remains assigned and completed for the impacted role set, and there is no recurrence signal based on the defined monitoring mechanism. The plan also defines the evidence set that will be reviewed at execution (training records, completed logs/work orders, configuration evidence, and a documented recurrence/trending assessment) and the expected documentation output (an executed EC review checklist/memo with a recorded conclusion). This EC should not be considered effective until the defined review is executed and a conclusion is recorded.
""".strip()


def build_system_prompt(record_id: str) -> str:
    rid = (record_id or "").upper()
    if rid.startswith("DEV-"):
        return "\n\n".join([BASE_PROMPT, DEV_ADDENDUM, DEV_EXAMPLE])
    if rid.startswith("EC-"):
        return "\n\n".join([BASE_PROMPT, EC_ADDENDUM, EC_EXAMPLE])
    if rid.startswith("CAPA-"):
        return "\n\n".join([BASE_PROMPT, CAPA_ADDENDUM, CAPA_EXAMPLE])
    if rid.startswith("CR-"):
        return "\n\n".join([BASE_PROMPT, CR_ADDENDUM, CR_EXAMPLE])
    return BASE_PROMPT


# -----------------------------
# Dataclasses
# -----------------------------

@dataclass
class ManifestRow:
    record: str
    record_id: str
    record_type: str
    file_path: str
    extracted_chars: int
    extract_attempts: int
    extract_method: str
    llm_ms: int
    title_chars: int
    synopsis_words: int
    status: str
    error_type: str
    error_message: str


@dataclass
class ErrorRow:
    record: str
    record_id: str
    record_type: str
    file_path: str
    error_type: str
    error_message: str
    attempts: int
    extracted_chars: int
    extract_method: str


# -----------------------------
# Helpers
# -----------------------------

def env_summary() -> str:
    return f"Python: {sys.version.split()[0]} | Executable: {sys.executable} | win32com={HAS_WIN32}"


def find_summary_docx_files(root_folder: str):
    matches = []
    for dirpath, _, filenames in os.walk(root_folder):
        for fn in filenames:
            if fn.lower().endswith("_summary.docx"):
                matches.append(os.path.join(dirpath, fn))
    return sorted(matches)


def record_stem_from_path(path: str) -> str:
    base = os.path.basename(path)
    stem, _ = os.path.splitext(base)
    return stem


def record_id_from_stem(stem: str) -> str:
    m = re.match(r"^([A-Z]{2,6}-\d{3,6}(?:-\d{1,3})?)_summary$", stem, flags=re.IGNORECASE)
    return m.group(1).upper() if m else stem.upper()


def record_type_from_id(record_id: str) -> str:
    rid = (record_id or "").upper()
    for prefix in ("CAPA-", "CR-", "DEV-", "EC-", "INV-", "OOS-"):
        if rid.startswith(prefix):
            return prefix[:-1]
    return rid.split("-", 1)[0] if "-" in rid else "UNKNOWN"


def load_api_key_from_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            s = line.strip()
            if s:
                return s
    return ""


def sanitize_title(title: str) -> str:
    title = re.sub(r"\s+", " ", (title or "").strip())
    if len(title) <= TITLE_CHAR_LIMIT:
        return title
    trunc = title[:TITLE_CHAR_LIMIT].rstrip()
    trunc2 = re.sub(r"\s+\S*$", "", trunc).rstrip()
    return trunc2 if trunc2 else trunc


def split_title_and_rest(text: str):
    t = (text or "").strip()
    if not t:
        return "", ""
    lines = t.splitlines()
    title = lines[0].strip()
    rest = "\n".join(lines[1:]).strip()
    if not rest and "\n\n" in t:
        a, b = t.split("\n\n", 1)
        return sanitize_title(a.strip()), b.strip()
    return sanitize_title(title), rest


def ensure_disclaimer_present(text: str) -> str:
    return text.strip() if "Legacy Migrated Record Disclaimer:" in text else (text.rstrip() + "\n\n" + DISCLAIMER_TEXT).strip()


def scrub_no_links(text: str) -> str:
    t = re.sub(r"https?://\S+", "", text or "")
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def word_count(s: str) -> int:
    return len(re.findall(r"\b\w+\b", s or ""))


# -----------------------------
# DOCX extraction
# -----------------------------

def docx_structure_counts(doc: Document):
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    header_para = header_tables = footer_para = footer_tables = 0
    for section in doc.sections:
        header_para += len(section.header.paragraphs)
        header_tables += len(section.header.tables)
        footer_para += len(section.footer.paragraphs)
        footer_tables += len(section.footer.tables)
    return {
        "para": para_count,
        "tables": table_count,
        "header_para": header_para,
        "header_tables": header_tables,
        "footer_para": footer_para,
        "footer_tables": footer_tables,
    }


def extract_docx_python(docx_path: str):
    doc = Document(docx_path)
    parts = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)

    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                t = (p.text or "").strip()
                if t:
                    parts.append(t)
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = (cell.text or "").strip()
                        if t:
                            parts.append(t)

    text = re.sub(r"\n{3,}", "\n\n", "\n".join(parts)).strip()
    counts = docx_structure_counts(doc)
    counts["extracted_chars"] = len(text)
    return text, counts


def extract_docx_tables_first(docx_path: str):
    doc = Document(docx_path)
    parts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    text = re.sub(r"\n{3,}", "\n\n", "\n".join(parts)).strip()
    counts = docx_structure_counts(doc)
    counts["extracted_chars"] = len(text)
    return text, counts


def normalize_windows_path(p: str) -> str:
    """Fixes OneDrive/forward-slash paths for Word COM."""
    ap = os.path.abspath(p)
    ap = ap.replace("/", "\\")
    return ap


def extract_docx_word_com(docx_path: str):
    if not HAS_WIN32:
        raise RuntimeError("pywin32/win32com not available. Install: python -m pip install pywin32")

    docx_path = normalize_windows_path(docx_path)
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(docx_path, ReadOnly=True)
        try:
            txt = doc.Content.Text or ""
        finally:
            doc.Close(False)
        txt = txt.replace("\r", "\n")
        return re.sub(r"\n{3,}", "\n\n", txt).strip()
    finally:
        word.Quit()


def classify_empty(struct_counts: dict) -> str:
    if any(struct_counts.get(k, 0) > 0 for k in ("tables", "para", "header_tables", "footer_tables", "header_para", "footer_para")):
        return "LIKELY_TEXTBOX_OR_IMAGE_CONTENT"
    return "EMPTY_DOCX_FINAL"


# -----------------------------
# OpenAI call
# -----------------------------

def call_openai(api_key: str, model: str, system_prompt: str, record_id: str, doc_text: str) -> str:
    client = OpenAI(api_key=api_key)
    user_input = f"RECORD_ID: {record_id}\n\nSOURCE_TEXT:\n{doc_text}"
    resp = client.responses.create(
        model=model,
        input=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_input},
        ],
    )
    out = getattr(resp, "output_text", None)
    if out:
        return out.strip()
    return str(resp)


# -----------------------------
# Tkinter GUI (real-time writer)
# -----------------------------

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("TrackWise DOCX _summary → CSV (REAL-TIME)")
        self.geometry("1180x820")

        self.in_root_var = tk.StringVar()
        self.out_folder_var = tk.StringVar()
        self.api_key_file_var = tk.StringVar()
        self.model_var = tk.StringVar(value=DEFAULT_MODEL)
        self.use_word_fallback_var = tk.BooleanVar(value=False)

        self.log_q = queue.Queue()
        self.stop_flag = threading.Event()
        self.worker = None

        self._build_ui()
        self._poll_log()
        self._log(env_summary())

    def _build_ui(self):
        pad = {"padx": 10, "pady": 8}
        root = ttk.Frame(self)
        root.pack(fill="both", expand=True)

        r1 = ttk.Frame(root); r1.pack(fill="x", **pad)
        ttk.Label(r1, text="Input folder (recursive):").pack(side="left")
        ttk.Entry(r1, textvariable=self.in_root_var, width=100).pack(side="left", padx=8, fill="x", expand=True)
        ttk.Button(r1, text="Browse…", command=self._browse_input).pack(side="left")

        r2 = ttk.Frame(root); r2.pack(fill="x", **pad)
        ttk.Label(r2, text="Output folder:").pack(side="left")
        ttk.Entry(r2, textvariable=self.out_folder_var, width=100).pack(side="left", padx=8, fill="x", expand=True)
        ttk.Button(r2, text="Browse…", command=self._browse_output).pack(side="left")

        r3 = ttk.Frame(root); r3.pack(fill="x", **pad)
        ttk.Label(r3, text="API key TXT file:").pack(side="left")
        ttk.Entry(r3, textvariable=self.api_key_file_var, width=100).pack(side="left", padx=8, fill="x", expand=True)
        ttk.Button(r3, text="Choose…", command=self._choose_key_file).pack(side="left")

        r4 = ttk.Frame(root); r4.pack(fill="x", **pad)
        ttk.Label(r4, text="Model:").pack(side="left")
        ttk.Entry(r4, textvariable=self.model_var, width=28).pack(side="left", padx=8)
        ttk.Label(r4, text=f"Title cap {TITLE_CHAR_LIMIT} | *_summary.docx").pack(side="left", padx=8)

        r4b = ttk.Frame(root); r4b.pack(fill="x", **pad)
        ttk.Checkbutton(
            r4b,
            text="Use Word fallback for empty docs (pywin32 + Word)",
            variable=self.use_word_fallback_var
        ).pack(side="left")

        r5 = ttk.Frame(root); r5.pack(fill="x", **pad)
        self.run_btn = ttk.Button(r5, text="Run", command=self._run)
        self.run_btn.pack(side="left")
        self.stop_btn = ttk.Button(r5, text="Stop", command=self._stop, state="disabled")
        self.stop_btn.pack(side="left", padx=8)

        self.progress = ttk.Progressbar(r5, orient="horizontal", mode="determinate", length=650)
        self.progress.pack(side="left", padx=12, fill="x", expand=True)

        self.status_lbl = ttk.Label(r5, text="Idle")
        self.status_lbl.pack(side="left")

        log_box = ttk.Labelframe(root, text="Event Log")
        log_box.pack(fill="both", expand=True, **pad)
        self.log_txt = tk.Text(log_box, wrap="word")
        self.log_txt.pack(fill="both", expand=True)
        self.log_txt.configure(state="disabled")

    def _browse_input(self):
        p = filedialog.askdirectory(title="Select input root folder")
        if p:
            self.in_root_var.set(p)

    def _browse_output(self):
        p = filedialog.askdirectory(title="Select output folder")
        if p:
            self.out_folder_var.set(p)

    def _choose_key_file(self):
        p = filedialog.askopenfilename(title="Select API key TXT", filetypes=[("Text files", "*.txt"), ("All files", "*.*")])
        if p:
            self.api_key_file_var.set(p)

    def _log(self, msg: str):
        self.log_q.put(msg)

    def _poll_log(self):
        try:
            while True:
                msg = self.log_q.get_nowait()
                self.log_txt.configure(state="normal")
                self.log_txt.insert("end", msg + "\n")
                self.log_txt.see("end")
                self.log_txt.configure(state="disabled")
        except queue.Empty:
            pass
        self.after(150, self._poll_log)

    def _set_running(self, running: bool):
        self.run_btn.configure(state="disabled" if running else "normal")
        self.stop_btn.configure(state="normal" if running else "disabled")
        self.status_lbl.configure(text="Running…" if running else "Idle")

    def _stop(self):
        self.stop_flag.set()
        self.status_lbl.configure(text="Stopping… (finalize after current record)")
        self._log("Stop requested: will stop after current record and finalize outputs (CSV/logs).")

    def _run(self):
        in_root = self.in_root_var.get().strip()
        out_folder = self.out_folder_var.get().strip()
        key_file = self.api_key_file_var.get().strip()
        model = self.model_var.get().strip() or DEFAULT_MODEL

        if not os.path.isdir(in_root):
            messagebox.showerror("Error", "Select a valid input folder.")
            return
        if not os.path.isdir(out_folder):
            messagebox.showerror("Error", "Select a valid output folder.")
            return
        if not os.path.isfile(key_file):
            messagebox.showerror("Error", "Select a valid API key TXT file.")
            return
        if OpenAI is None:
            messagebox.showerror("Missing dependency", f"Install openai:\n  {sys.executable} -m pip install --upgrade openai")
            return
        if self.use_word_fallback_var.get() and not HAS_WIN32:
            messagebox.showwarning("Word fallback not available", f"Install pywin32:\n  {sys.executable} -m pip install --upgrade pywin32")
            return

        api_key = load_api_key_from_txt(key_file)
        if not api_key:
            messagebox.showerror("Error", "API key TXT file is empty.")
            return

        files = find_summary_docx_files(in_root)
        if not files:
            messagebox.showwarning("No files", "No '*_summary.docx' files found under input folder.")
            return

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_csv = os.path.join(out_folder, f"record_title_synopsis_{ts}.csv")
        errors_csv = os.path.join(out_folder, f"errors_{ts}.csv")
        manifest_csv = os.path.join(out_folder, f"manifest_{ts}.csv")
        event_log_path = os.path.join(out_folder, f"event_log_{ts}.txt")

        self.stop_flag.clear()
        self.progress.configure(value=0, maximum=len(files))
        self._set_running(True)

        self._log(env_summary())
        self._log(f"MODE: REALTIME_WRITER=TRUE")
        self._log(f"Input root: {in_root}")
        self._log(f"Found {len(files)} summary DOCX files.")
        self._log(f"Model: {model}")
        self._log(f"Word fallback enabled: {self.use_word_fallback_var.get()}")
        self._log(f"Main CSV: {out_csv}")
        self._log(f"Errors CSV: {errors_csv}")
        self._log(f"Manifest CSV: {manifest_csv}")
        self._log(f"Event log: {event_log_path}")

        self.worker = threading.Thread(
            target=self._worker,
            args=(files, api_key, model, out_csv, errors_csv, manifest_csv, event_log_path),
            daemon=True
        )
        self.worker.start()

    def _worker(self, files, api_key, model, out_csv, errors_csv, manifest_csv, event_log_path):
        error_rows: list[ErrorRow] = []
        extract_chars_list: list[int] = []
        llm_ms_list: list[int] = []
        ok_count = 0
        fail_count = 0
        fail_by_type: dict[str, int] = {}
        processed = 0
        start_ts = time.time()

        manifest_header = list(asdict(ManifestRow(
            record="", record_id="", record_type="", file_path="",
            extracted_chars=0, extract_attempts=0, extract_method="",
            llm_ms=0, title_chars=0, synopsis_words=0,
            status="", error_type="", error_message=""
        )).keys())

        os.makedirs(os.path.dirname(out_csv) or ".", exist_ok=True)

        main_f = open(out_csv, "w", newline="", encoding="utf-8-sig")
        err_f = open(errors_csv, "w", newline="", encoding="utf-8-sig")
        man_f = open(manifest_csv, "w", newline="", encoding="utf-8-sig")
        log_f = open(event_log_path, "w", encoding="utf-8")

        main_w = csv.writer(main_f)
        err_w = csv.writer(err_f)
        man_w = csv.writer(man_f)

        main_w.writerow(["record", "title", "synopsis"])
        err_w.writerow(["record", "record_id", "record_type", "file_path", "error_type",
                        "error_message", "attempts", "extracted_chars", "extract_method"])
        man_w.writerow(manifest_header)

        def elog(line: str):
            self._log(line)
            log_f.write(line + "\n")

        def flush_all(reason: str):
            main_f.flush(); err_f.flush(); man_f.flush(); log_f.flush()
            elog(f"  FLUSH checkpoint: {reason} (rows_written={processed} errors_total={len(error_rows)})")

        try:
            for idx, path in enumerate(files, start=1):
                if self.stop_flag.is_set():
                    elog("STOPPING: stop flag detected. Exiting processing loop; finalizing outputs now…")
                    break

                stem = record_stem_from_path(path)         # record column must be e.g. DEV-00294_summary
                record_id = record_id_from_stem(stem)
                rtype = record_type_from_id(record_id)

                elog(f"[{idx}/{len(files)}] START record={stem} type={rtype} path={path}")

                # extraction retries
                extracted_text = ""
                extracted_chars = 0
                attempts = 0
                method = "python-docx"
                struct_counts_last = {}

                for attempt in range(1, 4):
                    attempts = attempt
                    t1, counts = extract_docx_python(path)
                    extracted_text = t1
                    extracted_chars = len(t1)
                    struct_counts_last = counts
                    elog(f"  EXTRACT attempt={attempt} method=python-docx chars={extracted_chars} "
                         f"para={counts.get('para')} tables={counts.get('tables')} "
                         f"hdr_para={counts.get('header_para')} hdr_tbl={counts.get('header_tables')} "
                         f"ftr_para={counts.get('footer_para')} ftr_tbl={counts.get('footer_tables')}")
                    if extracted_chars > 0:
                        break
                    time.sleep(0.25 * attempt)

                if extracted_chars == 0:
                    t2, counts2 = extract_docx_tables_first(path)
                    elog(f"  EXTRACT alt_pass method=tables-first chars={len(t2)} "
                         f"para={counts2.get('para')} tables={counts2.get('tables')}")
                    if len(t2) > 0:
                        extracted_text = t2
                        extracted_chars = len(t2)
                        struct_counts_last = counts2

                if extracted_chars == 0 and self.use_word_fallback_var.get():
                    try:
                        elog("  EXTRACT com_fallback attempting Word COM extraction…")
                        t3 = extract_docx_word_com(path)
                        if len(t3) > 0:
                            extracted_text = t3
                            extracted_chars = len(t3)
                            method = "word-com"
                            elog(f"  EXTRACT com_fallback SUCCESS chars={extracted_chars}")
                    except Exception as e:
                        elog(f"  EXTRACT com_fallback ERROR: {e}")

                if extracted_chars == 0:
                    error_type = classify_empty(struct_counts_last)
                    msg = "DOCX appears empty after extraction (paragraphs/tables/headers)."
                    elog(f"  COMPLETE status=ERROR error_type={error_type} message={msg}")

                    fail_count += 1
                    fail_by_type[error_type] = fail_by_type.get(error_type, 0) + 1

                    er = ErrorRow(stem, record_id, rtype, path, error_type, msg, attempts, 0, method)
                    error_rows.append(er)

                    main_w.writerow([stem, f"ERROR ({record_id})", f"ERROR: {msg}\n\n{DISCLAIMER_TEXT}"])
                    err_w.writerow([er.record, er.record_id, er.record_type, er.file_path, er.error_type,
                                    er.error_message, er.attempts, er.extracted_chars, er.extract_method])

                    mr = ManifestRow(stem, record_id, rtype, path, 0, attempts, method, 0, 0, 0, "ERROR", error_type, msg)
                    d = asdict(mr); man_w.writerow([d[k] for k in manifest_header])

                    processed += 1
                    elog(f"  WRITE row main_csv: {processed}/{len(files)}")
                    elog(f"  WRITE row errors_csv: +1 (total_errors={len(error_rows)})")
                    elog(f"  WRITE row manifest_csv: {processed}/{len(files)}")

                    self.after(0, lambda v=idx: self.progress.configure(value=v))
                    if processed % FLUSH_EVERY_N_RECORDS == 0:
                        flush_all("periodic")
                    continue

                extract_chars_list.append(extracted_chars)

                if self.stop_flag.is_set():
                    elog("STOPPING: stop flag set before LLM call. Finalizing outputs now…")
                    break

                # LLM
                system_prompt = build_system_prompt(record_id)
                llm_start = time.time()
                try:
                    elog(f"  LLM call model={model} input_chars={extracted_chars} prompt_type={rtype}")
                    llm_out = call_openai(api_key, model, system_prompt, record_id, extracted_text)
                    llm_ms = int((time.time() - llm_start) * 1000)
                    llm_ms_list.append(llm_ms)
                    elog(f"  LLM ok ms={llm_ms} out_chars={len(llm_out)}")
                except Exception as e:
                    llm_ms = int((time.time() - llm_start) * 1000)
                    error_type = "API_ERROR"
                    msg = str(e)
                    elog(f"  COMPLETE status=ERROR error_type={error_type} message={msg}")

                    fail_count += 1
                    fail_by_type[error_type] = fail_by_type.get(error_type, 0) + 1

                    er = ErrorRow(stem, record_id, rtype, path, error_type, msg, attempts, extracted_chars, method)
                    error_rows.append(er)

                    main_w.writerow([stem, f"ERROR ({record_id})", f"ERROR: {msg}\n\n{DISCLAIMER_TEXT}"])
                    err_w.writerow([er.record, er.record_id, er.record_type, er.file_path, er.error_type,
                                    er.error_message, er.attempts, er.extracted_chars, er.extract_method])

                    mr = ManifestRow(stem, record_id, rtype, path, extracted_chars, attempts, method, llm_ms, 0, 0, "ERROR", error_type, msg)
                    d = asdict(mr); man_w.writerow([d[k] for k in manifest_header])

                    processed += 1
                    elog(f"  WRITE row main_csv: {processed}/{len(files)}")
                    elog(f"  WRITE row errors_csv: +1 (total_errors={len(error_rows)})")
                    elog(f"  WRITE row manifest_csv: {processed}/{len(files)}")

                    self.after(0, lambda v=idx: self.progress.configure(value=v))
                    if processed % FLUSH_EVERY_N_RECORDS == 0:
                        flush_all("periodic")
                    continue

                # Parse
                title, rest = split_title_and_rest(llm_out)
                title = sanitize_title(title)
                rest = ensure_disclaimer_present(scrub_no_links(rest))

                synopsis_words = word_count(rest)
                elog(f"  PARSE ok title_chars={len(title)} synopsis_words={synopsis_words} disclaimer_present=Y")

                # Write rows immediately
                main_w.writerow([stem, title, rest])

                mr = ManifestRow(stem, record_id, rtype, path, extracted_chars, attempts, method, llm_ms,
                                 len(title), synopsis_words, "OK", "", "")
                d = asdict(mr); man_w.writerow([d[k] for k in manifest_header])

                processed += 1
                ok_count += 1
                elog("  COMPLETE status=OK row_written=Y")
                elog(f"  WRITE row main_csv: {processed}/{len(files)}")
                elog(f"  WRITE row manifest_csv: {processed}/{len(files)}")

                self.after(0, lambda v=idx: self.progress.configure(value=v))
                if processed % FLUSH_EVERY_N_RECORDS == 0:
                    flush_all("periodic")

            flush_all("final")

            elapsed_s = int(time.time() - start_ts)

            def fmt_stats(vals):
                if not vals:
                    return "n/a"
                return f"min={min(vals)} | median={int(statistics.median(vals))} | max={max(vals)}"

            elog("")
            elog("RUN SUMMARY")
            elog(f"  elapsed_s={elapsed_s}")
            elog(f"  total_discovered={len(files)}")
            elog(f"  processed={processed}")
            elog(f"  ok={ok_count}")
            elog(f"  failed={fail_count}")
            for k in sorted(fail_by_type.keys()):
                elog(f"  failed_{k}={fail_by_type[k]}")
            elog(f"  extracted_chars_stats: {fmt_stats(extract_chars_list)}")
            elog(f"  llm_ms_stats: {fmt_stats(llm_ms_list)}")
            elog("")
            elog(f"  output_csv={out_csv}")
            elog(f"  errors_csv={errors_csv}")
            elog(f"  manifest_csv={manifest_csv}")
            elog(f"  event_log={event_log_path}")

            elog("")
            elog("ERROR LIST (ALL)")
            if not error_rows:
                elog("  (none)")
            else:
                for i, er in enumerate(error_rows, start=1):
                    safe_msg = (er.error_message or "").replace("\n", " ").strip()
                    elog(
                        f"  {i:04d} | record={er.record} | type={er.record_type} | "
                        f"error_type={er.error_type} | attempts={er.attempts} | "
                        f"extracted_chars={er.extracted_chars} | method={er.extract_method} | "
                        f"path={er.file_path} | message={safe_msg}"
                    )

            elog("Done.")

            self.after(0, lambda: messagebox.showinfo(
                "Complete",
                f"Main CSV:\n{out_csv}\n\nErrors:\n{errors_csv}\nManifest:\n{manifest_csv}\nEvent log:\n{event_log_path}"
            ))

        finally:
            try: main_f.close()
            except Exception: pass
            try: err_f.close()
            except Exception: pass
            try: man_f.close()
            except Exception: pass
            try: log_f.close()
            except Exception: pass
            self.after(0, lambda: self._set_running(False))

    def _set_running(self, running: bool):
        self.run_btn.configure(state="disabled" if running else "normal")
        self.stop_btn.configure(state="normal" if running else "disabled")
        self.status_lbl.configure(text="Running…" if running else "Idle")


if __name__ == "__main__":
    App().mainloop()